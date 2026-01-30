"""Generate a Blogpost.docx from `docs/Blogpost.md`.

Usage:
  python scripts/generate_blog_docx.py --src docs/Blogpost.md --out docs/Blogpost.docx

This script uses `python-docx`. Install it with:
  pip install python-docx

It performs a simple markdown-to-docx conversion: '#' and '##' headings map to heading levels; paragraphs preserved; code fences are written verbatim using a monospace run.
"""
from pathlib import Path
import argparse


def md_to_docx(src_path: Path, out_path: Path):
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception as e:
        raise RuntimeError(
            "python-docx is required: pip install python-docx"
        ) from e

    doc = Document()

    text = src_path.read_text(encoding="utf-8")

    in_code = False
    for raw_line in text.splitlines():
        line = raw_line.rstrip('\n')
        if line.strip().startswith('```'):
            in_code = not in_code
            continue
        if in_code:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            continue

        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if not line.strip():
            doc.add_paragraph('')
            continue

        doc.add_paragraph(line)

    doc.save(str(out_path))


def main():
    p = argparse.ArgumentParser()
    p.add_argument('--src', default='docs/Blogpost.md')
    p.add_argument('--out', default='docs/Blogpost.docx')
    args = p.parse_args()

    src = Path(args.src)
    out = Path(args.out)

    if not src.exists():
        print('Source markdown not found:', src)
        raise SystemExit(1)

    md_to_docx(src, out)
    print('Wrote', out)


if __name__ == '__main__':
    main()
